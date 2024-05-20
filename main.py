import datetime
import os
import pickle
import random
import sys
import traceback
from io import BytesIO

import fitz
import numpy as np
import pygame
import requests
from PIL import Image
from PyQt6.QtCore import QThread, pyqtSignal, pyqtSlot
from PyQt6.QtWidgets import QFileDialog, QApplication, QMainWindow, QMessageBox, QDialog
from aip import AipOcr
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score

import cut
import erase
import feedback
import practice
import record
import settings
import sub_win
from win import Ui_MainWindow

all_ = {"math": [], "phy": [], "chem": [], "bio": []}


class TextRect:
    def __init__(self, inf, rate):
        self.text = inf["words"]
        self.flag = -1
        location = inf["location"]
        self.cut_rect = pygame.Rect(location["left"], location["top"], location["width"], location["height"])
        self.rect = pygame.Rect(int(self.cut_rect.x * rate), int(self.cut_rect.y * rate), int(self.cut_rect.w * rate),
                                int(self.cut_rect.h * rate))


class QuestionRect:
    def __init__(self, rect, rate, surface):
        self.cut_rect = [rect]
        self.rect = []
        self.touched = -1
        self.selected = False
        self.rate = rate
        self.surface = surface

    def draw(self):
        if self.touched > -1:
            color = (255, 0, 0)
        else:
            color = (0, 0, 255)
        for _i in self.rect:
            pygame.draw.rect(self.surface, color, _i, 1)
            if self.selected:
                surface = pygame.Surface((_i.w - 2, _i.h - 2), pygame.SRCALPHA)
                surface.fill((0, 255, 0, 100))
                self.surface.blit(surface, (_i.x + 1, _i.y + 1))

    def done(self):
        self.rect = []
        for _i in self.cut_rect:
            self.rect.append(pygame.Rect(int(_i.x * self.rate), int(_i.y * self.rate), int(_i.w * self.rate),
                                         int(_i.h * self.rate)))

    def sense(self):
        self.touched = -1
        for _i in range(len(self.rect)):
            if self.rect[_i].collidepoint(pygame.mouse.get_pos()):
                self.touched = _i


def pdf(pdf_path, image_path):
    pdf_doc = fitz.open(pdf_path)
    for pg in range(pdf_doc.page_count):
        pdf_page = pdf_doc[pg]
        zoom_x = 3.7
        zoom_y = 3.7
        mat = fitz.Matrix(zoom_x, zoom_y).prerotate(1)
        pix = pdf_page.get_pixmap(matrix=mat, alpha=False)
        if not os.path.exists(image_path):
            os.makedirs(image_path)
        if "__single__section__XZHRO" in pdf_path:
            last = "__single__section__XZHRO"
        else:
            last = ""
        pix._writeIMG(f"{image_path}/{pg}{last}.jpg", format_=7, jpg_quality=100)


class Page:
    def __init__(self, path, surface, pages, cut_, cut_sec):
        if os.path.splitext(path)[1] == ".pdf":
            pdf(path, "./__ache__")
            images = os.listdir("./__ache__")
            for p in images:
                Page("./__ache__" + "\\" + p, surface, pages, cut_, cut_sec)
            return
        fi = open(path, "rb")
        file_bin = fi.read()
        result = client.accurate(file_bin)["words_result"]
        image = pygame.image.load(BytesIO(file_bin))
        self.pil_image = Image.open(BytesIO(file_bin))
        fi.close()
        del file_bin
        size = image.get_size()
        rate = screen_size[0] / size[0]
        if size[1] * rate > screen_size[1]:
            rate = screen_size[1] / size[1]
        self.rate = rate
        self.image = pygame.transform.smoothscale(image,
                                                  (int(image.get_width() * rate), int(image.get_height() * rate)))
        self.rects = []
        pos = []
        for _i in result:
            if len(_i["words"]) > -1:
                rect = TextRect(_i, rate)
                pos.append([rect.cut_rect.center[0], 0])
                self.rects.append(rect)
        x = np.array(pos)
        if "__single__section__XZHRO" in path:
            self.section_num = 4
            k_means = KMeans(n_clusters=4).fit(x)
        elif not cut_ or cut_sec < 0:
            scores = []
            means = []
            for _k in range(2, 5, 1):
                k_means = KMeans(n_clusters=_k).fit(x)
                scores.append(silhouette_score(x, k_means.labels_, metric="euclidean"))
                means.append(k_means)
            self.section_num = scores.index(max(scores)) + 2
            k_means = means[self.section_num - 2]
        else:
            self.section_num = cut_sec
            k_means = KMeans(n_clusters=cut_sec).fit(x)
        centers = k_means.cluster_centers_.copy()
        labels = list(range(self.section_num))
        for _i in range(self.section_num - 1):
            for _j in range(self.section_num - 1):
                if centers[_j][0] > centers[_j + 1][0]:
                    labels[_j], labels[_j + 1] = labels[_j + 1], labels[_j]
                    centers[_j][0], centers[_j + 1][0] = centers[_j + 1][0], centers[_j][0]
        centers = k_means.cluster_centers_
        for _i in range(len(k_means.labels_)):
            flag = k_means.labels_[_i]
            if abs(self.rects[_i].cut_rect.center[0] - centers[flag][0]) < size[0] / self.section_num * 0.33 and \
                    self.rects[_i].cut_rect.w < size[0] / self.section_num:
                self.rects[_i].flag = labels.index(flag)
        for _i in range(len(self.rects) - 1):
            for _j in range(len(self.rects) - 1):
                if self.rects[_j].flag > self.rects[_j + 1].flag:
                    self.rects[_j], self.rects[_j + 1] = self.rects[_j + 1], self.rects[_j]
        section_rect = []
        last = -1
        for _i in self.rects:
            if _i.flag == -1:
                continue
            if _i.flag != last:
                last = _i.flag
                section_rect.append(_i.cut_rect.copy())
            if section_rect[last].x > _i.cut_rect.x:
                section_rect[last].w += section_rect[last].x - _i.cut_rect.x
                section_rect[last].x = _i.cut_rect.x
            if section_rect[last].w + section_rect[last].x < _i.cut_rect.w + _i.cut_rect.x:
                section_rect[last].w = _i.cut_rect.w + _i.cut_rect.x - section_rect[last].x
            section_rect[last].h = _i.cut_rect.y + _i.cut_rect.h - section_rect[last].y
        self.section_rects = section_rect
        if not cut_:
            num = (
            "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.", "13.", "14.", "15.", "16.",
            "17.", "18.", "19.", "20.", "21.", "22.")
            question_rects = []
            _item = -1
            section = -1
            rect_item = -1
            for _i in self.rects:
                if _i.flag == -1:
                    continue
                for _j in num:
                    if _i.text.startswith(_j):
                        if len(_i.text) != len(_j):
                            _i.text += "     "
                        if len(_i.text) == len(_j):
                            will_cut = True
                        elif not _i.text[len(_j)] in "0123456789":
                            will_cut = True
                        elif _i.text[len(_j) + 4] == "年" or _i.text[len(_j) + 1] == "分" or _i.text[len(_j) + 2] == "分" \
                                or _i.text[len(_j) + 3] == "分":
                            will_cut = True
                        else:
                            will_cut = False
                        if will_cut:
                            _item += 1
                            question_rects.append(QuestionRect(_i.cut_rect.copy(), rate, surface))
                            section = _i.flag
                            rect_item = 0
                            continue
                if _item == -1:
                    continue
                if _i.flag != section:
                    rect_item += 1
                    section = _i.flag
                    question_rects[_item].cut_rect.append(_i.cut_rect.copy())
                if question_rects[_item].cut_rect[rect_item].x > _i.cut_rect.x:
                    question_rects[_item].cut_rect[rect_item].w += question_rects[_item].cut_rect[
                                                                       rect_item].x - _i.cut_rect.x
                    question_rects[_item].cut_rect[rect_item].x = _i.cut_rect.x
                if question_rects[_item].cut_rect[rect_item].w + question_rects[_item].cut_rect[rect_item].x < \
                        _i.cut_rect.w + _i.cut_rect.x:
                    question_rects[_item].cut_rect[rect_item].w = _i.cut_rect.w + _i.cut_rect.x - \
                                                                  question_rects[_item].cut_rect[rect_item].x
                question_rects[_item].cut_rect[rect_item].h = _i.cut_rect.y + _i.cut_rect.h - \
                                                              question_rects[_item].cut_rect[rect_item].y
            for _i in question_rects:
                _i.done()
            self.question_rects = question_rects
        else:
            question_rects = []
            for i in self.section_rects:
                r = QuestionRect(i, rate, surface)
                r.done()
                r.selected = True
                question_rects.append(r)
            self.question_rects = question_rects
        pages.append(self)


def process(dir_path, sub, cut_, cut_sec, erase_handwriting, screen):
    sub_index = index[sub]
    pages = []
    item = 0
    paths = os.listdir(dir_path)
    file_item = 0
    pygame.display.set_caption("错题框选", "错题框选")
    for i in paths:
        file_item += 1
        path = dir_path + "\\" + i
        img = font.render(f"正在处理{path}", 1, (0, 0, 0))
        rect = img.get_rect()
        rect.center = screen.get_rect().center
        screen.fill((255, 255, 255))
        screen.blit(img, rect)
        pygame.display.update()
        Page(path, screen, pages, cut_, cut_sec)
    while True:
        for event in pygame.event.get():
            if event.type == pygame.QUIT:
                document = Document()
                docx_section = document.sections[0]
                if not cut_:
                    docx_section.page_width = Cm(18.2)
                    docx_section.page_height = Cm(25.7)
                    docx_section.top_margin = Cm(0.3)
                    docx_section.right_margin = Cm(0.9)
                    docx_section.bottom_margin = Cm(1.3)
                    docx_section.left_margin = Cm(1.4)
                    sect_pr = docx_section._sectPr
                    cols = sect_pr.xpath("./w:cols")[0]
                    cols.set(qn("w:num"), "2")
                    cols.set(qn("w:space"), "20")
                else:
                    docx_section.page_width = Cm(21)
                    docx_section.page_height = Cm(29.7)
                    docx_section.top_margin = Cm(0.3)
                    docx_section.right_margin = Cm(0.3)
                    docx_section.bottom_margin = Cm(0.3)
                    docx_section.left_margin = Cm(0.3)
                for k in pages:
                    for i in k.question_rects:
                        if i.selected:
                            img_size = [0, 0]
                            for j in i.cut_rect:
                                if j.w > img_size[0]:
                                    img_size[0] = j.w
                                img_size[1] += j.h
                            new = Image.new("RGB", (img_size[0], img_size[1]), (255, 255, 255))
                            y = 0
                            for j in i.cut_rect:
                                new.paste(k.pil_image.crop((j.x, j.y, j.x + j.w, j.y + j.h)), (0, y))
                                y += j.h
                            img_byte = BytesIO()
                            new.save(img_byte, format="JPEG", quality=100)
                            if not cut_:
                                num = len(sub_index)
                                with open(f"./{sub}/{num}.jpg", "wb") as file:
                                    info = {
                                        "path": f"./{sub}/{num}.jpg",
                                        "removed": f"./{sub}/{num}_removed.jpg",
                                        "err_num": 0,
                                        "processed": False,
                                        "date": datetime.date.today(),
                                        "practiced": False,
                                        "pra_num": 0,
                                        "practiced_date": None,
                                        "correct": None,
                                        "short_ans": "",
                                        "detailed_ans": "",
                                        "printed": False,
                                        "point": ""
                                    }
                                    sub_index.append(info)
                                    file.write(img_byte.getvalue())
                                if options["erase"]:
                                    font_img = font.render(f"正在为{num}.jpg去手写", 1, (0, 0, 0))
                                    rect = font_img.get_rect()
                                    rect.center = screen.get_rect().center
                                    screen.fill((255, 255, 255))
                                    screen.blit(font_img, rect)
                                    pygame.display.update()
                                    with open(f"./{sub}/{num}_removed.jpg", "wb") as file:
                                        img_byte.seek(0)
                                        file.write(erase.create_request(img_byte))
                                document.add_picture(img_byte, width=Cm(7.7))
                            else:
                                if options["erase"] and erase_handwriting:
                                    img_byte.seek(0)
                                    er = BytesIO(erase.create_request(img_byte))
                                else:
                                    er = img_byte
                                if 20.4 / img_size[0] * img_size[1] > 29.1:
                                    document.add_picture(er, height=Cm(29.1))
                                else:
                                    document.add_picture(er, width=Cm(20.4))
                if not cut_:
                    with open(f"./{sub}/index.pkl", "wb") as file:
                        pickle.dump(sub_index, file)
                return document
            if event.type == pygame.MOUSEBUTTONDOWN:
                for r in pages[item].question_rects:
                    if r.touched > -1:
                        r.selected = not r.selected
            if event.type == pygame.KEYDOWN:
                if event.key == pygame.K_RIGHT:
                    item += 1
                    if item >= len(pages):
                        item = 0
                if event.key == pygame.K_LEFT:
                    item -= 1
                    if item < 0:
                        item = len(pages) - 1
                if event.key == pygame.K_DELETE:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            del r.cut_rect[r.touched]
                            del r.rect[r.touched]
                            if len(r.cut_rect) == 0:
                                r.selected = False
                if event.key == pygame.K_a:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].x -= 10
                            r.cut_rect[r.touched].w += 10
                            r.done()
                if event.key == pygame.K_w:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].y -= 10
                            r.cut_rect[r.touched].h += 10
                            r.done()
                if event.key == pygame.K_d:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].w += 10
                            r.done()
                if event.key == pygame.K_s:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].h += 10
                            r.done()
                if event.key == pygame.K_l:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].x += 10
                            r.cut_rect[r.touched].w -= 10
                            r.done()
                if event.key == pygame.K_k:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].y += 10
                            r.cut_rect[r.touched].h -= 10
                            r.done()
                if event.key == pygame.K_j:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].w -= 10
                            r.done()
                if event.key == pygame.K_i:
                    for r in pages[item].question_rects:
                        if r.touched > -1:
                            r.cut_rect[r.touched].h -= 10
                            r.done()
        screen.fill((255, 255, 255))
        screen.blit(pages[item].image, (0, 0))
        for r in pages[item].question_rects:
            r.sense()
            r.draw()
        pygame.display.update()


def err_handler(path, sub, cut_, cut_sec, erase_handwriting):
    screen = pygame.display.set_mode(screen_size)
    try:
        val = process(path, sub, cut_, cut_sec, erase_handwriting, screen)
        pygame.quit()
        return val
    except:
        pygame.quit()
        err = traceback.format_exc()
        feedback.send(err)
        return err


def practice_mistakes(sub, method, num):
    win.setEnabled(False)
    idx = index[sub]
    questions = []
    ids = []
    heights = []
    big = []
    big_ids = []
    big_heights = []
    sign = {"math": "MA", "phy": "PH", "chem": "CH", "bio": "BI"}[sub]
    if method == 0:
        while len(questions) + len(big) < num:
            random_question_idx = random.randint(0, len(idx) - 1)
            random_question = idx[random_question_idx]
            if random_question not in questions and random_question not in big:
                all_[sub].append(random_question_idx)
                img = Image.open(random_question["removed"])
                idx[random_question_idx]["practiced"] = True
                idx[random_question_idx]["practiced_date"] = datetime.date.today()
                if img.size[0] < 1650:
                    questions.append(random_question)
                    ids.append(idx.index(random_question))
                    heights.append(img.size[1] / img.size[0])
                else:
                    big.append(random_question)
                    big_ids.append(idx.index(random_question))
                    big_heights.append(img.size[1] / img.size[0])
    for j in range(len(questions) - 1):
        for i in range(len(questions) - 1):
            if heights[i] > heights[i + 1]:
                heights[i], heights[i + 1] = heights[i + 1], heights[i]
                questions[i], questions[i + 1] = questions[i + 1], questions[i]
                ids[i], ids[i + 1] = ids[i + 1], ids[i]
    for j in range(len(big) - 1):
        for i in range(len(big) - 1):
            if big_heights[i] > big_heights[i + 1]:
                big_heights[i], big_heights[i + 1] = big_heights[i + 1], big_heights[i]
                big[i], big[i + 1] = big[i + 1], big[i]
                big_ids[i], big_ids[i + 1] = big_ids[i + 1], big_ids[i]
    document = Document()
    docx_section = document.sections[0]
    docx_section.page_width = Cm(21)
    docx_section.page_height = Cm(29.7)
    docx_section.top_margin = Cm(1)
    docx_section.right_margin = Cm(1)
    docx_section.bottom_margin = Cm(1)
    docx_section.left_margin = Cm(1)
    sect_pr = docx_section._sectPr
    cols = sect_pr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "20")
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(8)
    para = document.add_paragraph("错题重练")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].font.size = Pt(18)
    para.runs[0].font.bold = True
    for i in range(len(questions)):
        document.add_picture(questions[i]["removed"], width=Cm(9.5))
        p = document.add_paragraph(f"  试题编号：{sign}{ids[i]}")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        p.line_spacing = 1
    new_section = document.add_section()
    sect_pr = new_section._sectPr
    cols = sect_pr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "1")
    cols.set(qn("w:space"), "1")
    for i in range(len(big)):
        document.add_picture(big[i]["removed"], width=Cm(19))
        p = document.add_paragraph(f"  试题编号：{sign}{big_ids[i]}")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    ending = document.add_paragraph("练习结束，记得将错误情况录入错题整理程序")
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    win.setEnabled(True)
    document.save(QFileDialog.getSaveFileName(win, "保存", "C:/", "Word 文档 (*.docx)")[0])
    win.record_mistakes()
    with open(f"./{sub}/index.pkl", "wb") as f_:
        pickle.dump(idx, f_)


def first_page(sub):
    win.setEnabled(False)
    idx = index[sub]
    questions = []
    ids = []
    quotient = []
    big = []
    big_ids = []
    big_quotient = []
    sign = {"math": "MA", "phy": "PH", "chem": "CH", "bio": "BI"}[sub]
    for question_idx in range(len(idx)):
        question = idx[question_idx]
        if not question["printed"]:
            img = Image.open(question["path"])
            # idx[question_idx]["printed"] = True
            if img.size[0] < 1650:
                questions.append(question)
                ids.append(question_idx)
                quotient.append(img.size[1] / img.size[0])
            else:
                big.append(question)
                big_ids.append(question_idx)
                big_quotient.append(img.size[1] / img.size[0])
    with open(f"./{sub}/index.pkl", "wb") as f_:
        pickle.dump(idx, f_)
    for j in range(len(questions) - 1):
        for i in range(len(questions) - 1):
            if quotient[i] < quotient[i + 1]:
                quotient[i], quotient[i + 1] = quotient[i + 1], quotient[i]
                questions[i], questions[i + 1] = questions[i + 1], questions[i]
                ids[i], ids[i + 1] = ids[i + 1], ids[i]
    for j in range(len(big) - 1):
        for i in range(len(big) - 1):
            if big_quotient[i] < big_quotient[i + 1]:
                big_quotient[i], big_quotient[i + 1] = big_quotient[i + 1], big_quotient[i]
                big[i], big[i + 1] = big[i + 1], big[i]
                big_ids[i], big_ids[i + 1] = big_ids[i + 1], big_ids[i]
    j = 0
    q_sum = quotient[0]
    di = 1
    while not quotient[j] == quotient[-1]:
        q_sum += quotient[-1]
        if q_sum > 3.01:
            j += di
            di = 1
            if j > len(quotient) - 1:
                break
            q_sum = quotient[j]
            continue
        quotient.insert(j, quotient.pop())
        questions.insert(j, questions.pop())
        ids.insert(j, ids.pop())
        di += 1
    if len(big) > 0:
        j = 0
        q_sum = big_quotient[0] + 0.21135469
        di = 1
        while not big_quotient[j] == big_quotient[-1]:
            q_sum += big_quotient[-1] + 0.21135469
            if q_sum > 1.486:
                j += di
                di = 1
                if j > len(big_quotient) - 1:
                    break
                q_sum = big_quotient[j] + 0.21135469
                continue
            big_quotient.insert(j, big_quotient.pop())
            big.insert(j, big.pop())
            big_ids.insert(j, big_ids.pop())
            di += 1
    document = Document()
    docx_section = document.sections[0]
    docx_section.page_width = Cm(18.2)
    docx_section.page_height = Cm(25.7)
    docx_section.top_margin = Cm(0.3)
    docx_section.right_margin = Cm(0.9)
    docx_section.bottom_margin = Cm(1.3)
    docx_section.left_margin = Cm(1.4)
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(8)
    q_sum = 0
    for i in range(len(questions)):
        q_sum += quotient[i]
        if q_sum > 3.01:
            q_sum = quotient[i]
            document.add_section()
        table = document.add_table(3, 2, style="Table Grid")
        cell_new = table.cell(0, 0).merge(table.cell(1, 0)).merge(table.cell(2, 0))
        cell_new.paragraphs[0].add_run().add_picture(questions[i]["path"], width=Cm(8))
        table.cell(0, 1).text = f"试题编号：{sign}{ids[i]}"
        table.cell(1, 1).text = "知识点："
        table.cell(2, 1).text = "笔记："
    if len(big) > 0:
        document.add_section()
        q_sum = 0
        for i in range(len(big)):
            q_sum += big_quotient[i] + 0.21135469
            if q_sum > 1.486:
                q_sum = big_quotient[i] + 0.21135469
                document.add_section()
            table = document.add_table(3, 2, style="Table Grid")
            cell_new = table.cell(0, 0).merge(table.cell(0, 1))
            cell_new.paragraphs[0].add_run().add_picture(big[i]["path"], width=Cm(15.7))
            table.cell(1, 0).text = f"试题编号：{sign}{big_ids[i]}"
            table.cell(1, 1).text = "知识点："
            new = table.cell(2, 0).merge(table.cell(2, 1))
            new.text = "笔记：\n\n\n\n\n\n\n"
    win.setEnabled(True)
    document.save(QFileDialog.getSaveFileName(win, "保存", "C:/", "Word 文档 (*.docx)")[0])


class Thread(QThread):
    signal_tuple = pyqtSignal(tuple)

    def __init__(self, func):
        super().__init__()
        self.func = func

    def run(self):
        result = self.func()
        self.signal_tuple.emit((result, 1))


def practice_starter(sub, num, method, ignore, ignore_range, range_e):
    sign = {"数学": "math", "物理": "phy", "化学": "chem", "生物": "bio"}
    practice_mistakes(sign[sub], 0, num)


def process_mistakes(text):
    global all_
    err_list = text.upper().split("\n")
    sign = {"MA": "math", "PH": "phy", "CH": "chem", "BI": "bio"}
    errs = {"math": [], "phy": [], "chem": [], "bio": []}
    for i in err_list:
        if len(i) > 2:
            errs[sign[i[:2]]].append(int(i[2:]))
    for key in errs.keys():
        for i in all_[key]:
            index[key][i]["pra_num"] += 1
            if i in errs[key]:
                index[key][i]["correct"] = False
                index[key][i]["err_num"] += 1
            else:
                index[key][i]["correct"] = True
    all_ = {"math": [], "phy": [], "chem": [], "bio": []}


def save_practiced():
    global all_
    with open("./practiced.pkl", "wb") as pr:
        pickle.dump(all_, pr)
    all_ = {"math": [], "phy": [], "chem": [], "bio": []}


class Window(QMainWindow, Ui_MainWindow):
    def __init__(self, parent=None):
        QMainWindow.__init__(self, parent=parent)
        self.setupUi(self)
        self.actionOpen.triggered.connect(lambda: self.choose_sub(self.normal_open))
        self.pushButtonMath.clicked.connect(lambda: self.setup_thread("math", False, -1, False))
        self.pushButtonPhysics.clicked.connect(lambda: self.setup_thread("phy", False, -1, False))
        self.pushButtonChem.clicked.connect(lambda: self.setup_thread("chem", False, -1, False))
        self.pushButtonBio.clicked.connect(lambda: self.setup_thread("bio", False, -1, False))
        self.actionSettings.triggered.connect(self.settings)
        self.actionCut_Paper.triggered.connect(self.cut_paper)
        self.actionPraErr.triggered.connect(self.practice_win)
        self.actionExport_Docx.triggered.connect(lambda: self.choose_sub(first_page))
        self.actionExport_Mistakes.setEnabled(False)
        self.actionProcess_Mistakes.setEnabled(False)
        self.actionSave.setEnabled(False)
        self.actionContact_Us.setEnabled(False)
        self.actionAbout.setEnabled(False)
        self.actionCheck_Updates.triggered.connect(self.check_updates)

    def setup_thread(self, sub, cut_, cut_sec, erase_handwriting):
        dir_path = QFileDialog.getExistingDirectory(win, "浏览", "C:/")
        self.setEnabled(False)
        self.thread_ = Thread(lambda: err_handler(dir_path, sub, cut_, cut_sec, erase_handwriting))
        self.thread_.signal_tuple.connect(self.thread_finished)
        self.thread_.start()

    def move_center(self, obj):
        obj.move(int((self.pos().x() + win.size().width() / 2 - obj.size().width() / 2)),
                 int(self.pos().y() + win.size().height() / 2 - obj.size().height() / 2))

    @pyqtSlot(tuple)
    def thread_finished(self, item):
        self.setEnabled(True)
        if isinstance(item[0], str):
            QMessageBox.critical(win, "程序运行出错", f"错误信息：\n{item[0]}")
        else:
            item[0].save(QFileDialog.getSaveFileName(win, "保存", "C:/", "Word 文档 (*.docx)")[0])

    def settings(self):
        class Settings(QDialog, settings.Ui_Dialog):
            def __init__(self, parent=win):
                QDialog.__init__(self, parent)
                self.setupUi(self)
                self.checkBox.setEnabled(False)
                self.lineEdit.setText(str(options["size"][0]))
                self.lineEdit_2.setText(str(options["size"][1]))
                self.checkBox.setChecked(options["online"])
                self.lineEdit_3.setText(options["APP_ID"])
                self.lineEdit_4.setText(options["API_KEY"])
                self.lineEdit_5.setText(options["SECRET_KEY"])
                self.checkBox_2.setChecked(options["erase"])
                self.lineEdit_6.setText(options["yd_id"])
                self.lineEdit_7.setText(options["yd_key"])

        settings_window = Settings()
        self.move_center(settings_window)
        settings_window.exec()
        global options
        options = {
            "size": (int(settings_window.lineEdit.text()), int(settings_window.lineEdit_2.text())),
            "online": settings_window.checkBox.isChecked(),
            "APP_ID": settings_window.lineEdit_3.text(),
            "API_KEY": settings_window.lineEdit_4.text(),
            "SECRET_KEY": settings_window.lineEdit_5.text(),
            "erase": settings_window.checkBox_2.isChecked(),
            "yd_id": settings_window.lineEdit_6.text(),
            "yd_key": settings_window.lineEdit_7.text(),
        }
        with open("./options.pkl", "wb") as f_:
            pickle.dump(options, f_)

    def cut_starter(self, text, erase_handwriting):
        if text == "自动识别":
            num = -1
        else:
            num = int(text)
        self.setup_thread("math", True, num, erase_handwriting)

    def cut_paper(self):
        class Cut(QDialog, cut.Ui_Dialog):
            def __init__(self, parent=win):
                QDialog.__init__(self, parent)
                self.setupUi(self)

        cut_win = Cut()
        self.move_center(cut_win)
        cut_win.accepted.connect(lambda: self.cut_starter(cut_win.comboBox.currentText(), cut_win.checkBox.isChecked()))
        cut_win.exec()

    def choose_sub(self, func):
        class Sub(QDialog, sub_win.Ui_Dialog):
            def __init__(self, parent=win):
                QDialog.__init__(self, parent)
                self.setupUi(self)
                sign = {"数学": "math", "物理": "phy", "化学": "chem", "生物": "bio"}
                self.accepted.connect(lambda: func(sign[self.comboBox.currentText()]))

        sub = Sub()
        self.move_center(sub)
        sub.exec()

    def normal_open(self, sub):
        self.setup_thread(sub, False, -1, False)

    def practice_win(self):
        class Practice(QDialog, practice.Ui_Dialog):
            def __init__(self, parent=self):
                QDialog.__init__(self, parent)
                self.setupUi(self)
                self.comboBox.setEnabled(False)
                self.comboBox_2.setEnabled(False)
                self.comboBox_4.setEnabled(False)
                self.checkBox.setEnabled(False)

        pw = Practice()
        self.move_center(pw)
        pw.accepted.connect(lambda: practice_starter(pw.comboBox_3.currentText(),
                                                     int(pw.spinBox.text()),
                                                     pw.comboBox.currentText(),
                                                     pw.checkBox.isChecked(),
                                                     pw.comboBox_2.currentText(),
                                                     pw.comboBox_4.currentText()))
        pw.exec()

    def record_mistakes(self):
        class Record(QDialog, record.Ui_Dialog):
            def __init__(self, parent=self):
                QDialog.__init__(self, parent)
                self.setupUi(self)
                self.accepted.connect(lambda: process_mistakes(self.textEdit.toPlainText()))
                self.rejected.connect(save_practiced)

        rec = Record()
        self.move_center(rec)
        rec.exec()

    def check_updates(self):
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                                 "Chrome/125.0.0.0 Safari/537.36 Edg/125.0.0.0"}
        response = requests.get("https://raw.githubusercontent.com/xzhiter/Mistakes_Recorder/main/version",
                                headers=headers, timeout=5)
        ver = float(response.text)
        if ver > 3.1:
            QMessageBox.information(self, "软件更新", f"检测到最新版本V{ver}")
        else:
            QMessageBox.information(self, "软件更新", "当前已是最新版本")


def except_hook(cls, exception, _traceback):
    err = "".join(traceback.format_exception(cls, exception, _traceback))
    feedback.send(err)
    QMessageBox.critical(None, "程序运行出错", f"错误信息：\n{err}")


if __name__ == "__main__":
    if os.path.isfile("./options.pkl"):
        with open("./options.pkl", "rb") as f:
            options = pickle.load(f)
    else:
        options = {
            "size": (1500, 900),
            "online": True,
            "APP_ID": "",
            "API_KEY": "",
            "SECRET_KEY": "",
            "erase": False,
            "yd_id": "",
            "yd_key": ""
        }
        with open("./options.pkl", "wb") as f:
            pickle.dump(options, f)
    if not os.path.isdir("./math"):
        os.mkdir("./math")
    if os.path.isfile("./math/index.pkl"):
        with open("./math/index.pkl", "rb") as f:
            math_index = pickle.load(f)
    else:
        math_index = []
        with open("./math/index.pkl", "wb") as f:
            pickle.dump(math_index, f)
    if not os.path.isdir("./phy"):
        os.mkdir("./phy")
    if os.path.isfile("./phy/index.pkl"):
        with open("./phy/index.pkl", "rb") as f:
            phy_index = pickle.load(f)
    else:
        phy_index = []
        with open("./phy/index.pkl", "wb") as f:
            pickle.dump(phy_index, f)
    if not os.path.isdir("./chem"):
        os.mkdir("./chem")
    if os.path.isfile("./chem/index.pkl"):
        with open("./chem/index.pkl", "rb") as f:
            chem_index = pickle.load(f)
    else:
        chem_index = []
        with open("./chem/index.pkl", "wb") as f:
            pickle.dump(chem_index, f)
    if not os.path.isdir("./bio"):
        os.mkdir("./bio")
    if os.path.isfile("./bio/index.pkl"):
        with open("./bio/index.pkl", "rb") as f:
            bio_index = pickle.load(f)
    else:
        bio_index = []
        with open("./bio/index.pkl", "wb") as f:
            pickle.dump(bio_index, f)
    index = {
        "chem": chem_index,
        "phy": phy_index,
        "math": math_index,
        "bio": bio_index
    }
    screen_size = options["size"]
    client = AipOcr(options["APP_ID"], options["API_KEY"], options["SECRET_KEY"])
    if options["erase"]:
        erase.set_key(options["yd_id"], options["yd_key"])
    pygame.init()
    font = pygame.font.Font("./msyh.ttc", 20)
    app = QApplication(sys.argv)
    win = Window()
    win.show()
    sys.excepthook = except_hook
    if os.path.isfile("./practiced.pkl"):
        with open("./practiced.pkl", "rb") as f:
            all_ = pickle.load(f)
        os.remove("./practiced.pkl")
        win.record_mistakes()
    app.exec()
